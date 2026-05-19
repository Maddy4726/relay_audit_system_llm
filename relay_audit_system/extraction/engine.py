"""
OpenAI Responses API extraction engine for relay feeder test reports.

Reads DOCX source documents, extracts structured data against the canonical
Pydantic schema, and writes JSON to ``extracted_json/``.
"""

from __future__ import annotations

import json
import logging
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from openai import APIConnectionError, APIStatusError, APITimeoutError, OpenAI, RateLimitError
from pydantic import ValidationError

from relay_audit_system.extraction.exceptions import (
    ConfigurationError,
    DocxLoadError,
    ExtractionError,
)
from relay_audit_system.schemas import ExtractionMetadata, RelayFeederTestReport
from relay_audit_system.utils.config import Settings, load_settings

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# DOCX loading
# ---------------------------------------------------------------------------


def load_docx_text(docx_path: Path) -> str:
    """
    Load plain text from a DOCX file, including paragraph and table content.

    Args:
        docx_path: Path to the ``.docx`` file.

    Returns:
        Normalized document text with paragraphs and table rows separated by newlines.

    Raises:
        DocxLoadError: If the file is missing, not a valid DOCX, or yields no text.
    """
    path = Path(docx_path)
    if not path.is_file():
        raise DocxLoadError(f"DOCX file not found: {path}")
    if path.suffix.lower() != ".docx":
        raise DocxLoadError(f"Expected a .docx file, got: {path.name}")

    try:
        document = Document(path)
    except PackageNotFoundError as exc:
        raise DocxLoadError(f"Invalid or corrupt DOCX package: {path}") from exc
    except OSError as exc:
        raise DocxLoadError(f"Unable to read DOCX file: {path}") from exc
    except Exception as exc:
        raise DocxLoadError(f"Unexpected error reading DOCX: {path}") from exc

    sections: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            sections.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        sections.append(f"[TABLE {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                sections.append(" | ".join(cells))

    combined = "\n".join(sections).strip()
    if not combined:
        raise DocxLoadError(f"DOCX contains no extractable text: {path}")

    logger.info("Loaded DOCX text (%d characters) from %s", len(combined), path.name)
    return combined


# ---------------------------------------------------------------------------
# Prompt
# ---------------------------------------------------------------------------


def build_extraction_prompt() -> str:
    """
    Build system instructions for structured extraction.

    The model must extract only what is explicitly present in the document.
    """
    return """You are a precision document extraction system for industrial relay \
feeder testing reports used in power plants and substations.

Your task is EXTRACTION ONLY. Convert the supplied report text into the provided \
JSON schema. You are not an auditor, reviewer, or engineer.

STRICT RULES:
1. EXTRACT ONLY — Copy values that appear in the document. Do not audit, judge, \
comment on pass/fail adequacy, or recommend actions.
2. NEVER INFER — Do not guess, calculate, derive, or fill gaps from domain knowledge. \
If a value is not explicitly stated, use null.
3. NULL FOR MISSING — Use null for unavailable fields. Use empty lists [] only when \
the document clearly has no rows for a repeatable section; otherwise include each \
row found.
4. PRESERVE UNITS — Keep units exactly as written (e.g. "A", "kV", "MVA", "s", "%", \
"µΩ"). Store values with units in the same form as the report when ambiguous.
5. PRESERVE PROTECTION CODES — Keep protection element codes verbatim (e.g. "50", \
"51", "49", "46", "DTOC", "IDMT", "REF", "87T"). Do not rename or normalize codes.
6. TABLES — Map table rows to the appropriate list fields (CTs, relays, protection \
tests, test points, measurements).
7. MULTIPLE RELAYS / ELEMENTS — Emit one list item per relay, CT, or protection test \
block when the report lists them separately.
8. FEEDER TYPES — Set feeder_type only when the document identifies motor feeder, \
transformer feeder, or bus coupler. Use motor_details, transformer_details, or \
bus_coupler_details only when relevant content exists.
9. RELAY TECHNOLOGY — Use numerical_details or electromechanical_details only when \
the report provides that class of data.
10. NO HALLUCINATION — Do not invent names, dates, serial numbers, settings, or test \
results.

Populate extraction_metadata with source_file and extraction warnings if text is \
illegible or truncated; do not fabricate data to compensate."""


# ---------------------------------------------------------------------------
# OpenAI structured extraction
# ---------------------------------------------------------------------------


def _build_user_message(document_text: str, *, source_file: str) -> str:
    return (
        f"Source file: {source_file}\n\n"
        "Extract all relay feeder test report data from the document below "
        "into the required JSON schema.\n\n"
        "--- DOCUMENT START ---\n"
        f"{document_text}\n"
        "--- DOCUMENT END ---"
    )


def _attach_extraction_metadata(
    report: RelayFeederTestReport,
    *,
    source_file: str,
    model_name: str,
    warnings: list[str] | None = None,
) -> RelayFeederTestReport:
    """Merge pipeline provenance into ``extraction_metadata``."""
    existing = report.extraction_metadata
    merged_warnings = list(warnings or [])
    if existing and existing.warnings:
        merged_warnings = [*existing.warnings, *merged_warnings]

    metadata = ExtractionMetadata(
        source_file=source_file,
        extracted_at=datetime.now(timezone.utc).isoformat(),
        extractor_version="relay_audit_system",
        model_name=model_name,
        warnings=merged_warnings,
        errors=list(existing.errors) if existing and existing.errors else [],
        field_confidence=dict(existing.field_confidence) if existing and existing.field_confidence else {},
        confidence_score=existing.confidence_score if existing else None,
        source_page_count=existing.source_page_count if existing else None,
        raw_blocks=list(existing.raw_blocks) if existing and existing.raw_blocks else [],
    )
    return report.model_copy(update={"extraction_metadata": metadata})


def extract_structured_json(
    document_text: str,
    *,
    source_file: str,
    settings: Settings | None = None,
    client: OpenAI | None = None,
) -> RelayFeederTestReport:
    """
    Call the OpenAI Responses API with structured Pydantic output.

    Args:
        document_text: Full text from ``load_docx_text()``.
        source_file: Original filename for metadata and logging.
        settings: Optional settings; loaded from environment when omitted.
        client: Optional pre-configured OpenAI client (for testing).

    Returns:
        Validated ``RelayFeederTestReport`` instance.

    Raises:
        ConfigurationError: If settings cannot be loaded.
        ExtractionError: On API, timeout, rate limit, or parse failures.
    """
    if not document_text.strip():
        raise ExtractionError("Document text is empty; cannot extract.")

    try:
        cfg = settings or load_settings()
    except ConfigurationError:
        raise
    except Exception as exc:
        raise ConfigurationError(f"Failed to load settings: {exc}") from exc

    openai_client = client or OpenAI(api_key=cfg.openai_api_key)
    instructions = build_extraction_prompt()
    user_message = _build_user_message(document_text, source_file=source_file)

    logger.info(
        "Calling OpenAI Responses API (model=%s, source=%s, chars=%d)",
        cfg.openai_model,
        source_file,
        len(document_text),
    )

    try:
        response = openai_client.responses.parse(
            model=cfg.openai_model,
            instructions=instructions,
            input=user_message,
            text_format=RelayFeederTestReport,
        )
    except RateLimitError as exc:
        raise ExtractionError("OpenAI rate limit exceeded.") from exc
    except APITimeoutError as exc:
        raise ExtractionError("OpenAI request timed out.") from exc
    except APIConnectionError as exc:
        raise ExtractionError("Unable to connect to OpenAI API.") from exc
    except APIStatusError as exc:
        raise ExtractionError(
            f"OpenAI API error (status={exc.status_code}): {exc.message}"
        ) from exc
    except Exception as exc:
        raise ExtractionError(f"OpenAI request failed: {exc}") from exc

    if getattr(response, "error", None):
        raise ExtractionError(f"OpenAI response error: {response.error}")

    report = response.output_parsed
    if report is None:
        raise ExtractionError(
            "OpenAI returned no parsed structured output. "
            "The model may have refused or truncated the response."
        )

    try:
        validated = RelayFeederTestReport.model_validate(report)
    except ValidationError as exc:
        raise ExtractionError(f"Structured output failed schema validation: {exc}") from exc

    warnings: list[str] = []
    if getattr(response, "status", None) == "incomplete":
        warnings.append("OpenAI response status was incomplete; some fields may be missing.")

    result = _attach_extraction_metadata(
        validated,
        source_file=source_file,
        model_name=cfg.openai_model,
        warnings=warnings or None,
    )
    logger.info("Structured extraction completed for %s", source_file)
    return result


# ---------------------------------------------------------------------------
# JSON persistence
# ---------------------------------------------------------------------------


def save_json(
    report: RelayFeederTestReport,
    output_path: Path,
    *,
    indent: int = 2,
) -> Path:
    """
    Serialize a report to JSON and write to disk.

    Args:
        report: Validated extraction result.
        output_path: Destination ``.json`` file path.
        indent: JSON indentation level.

    Returns:
        Resolved output path.

    Raises:
        ExtractionError: If the file cannot be written.
    """
    path = Path(output_path)
    if path.suffix.lower() != ".json":
        path = path.with_suffix(".json")

    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        payload: dict[str, Any] = report.model_dump(mode="json", exclude_none=False)
        path.write_text(
            json.dumps(payload, indent=indent, ensure_ascii=False) + "\n",
            encoding="utf-8",
        )
    except OSError as exc:
        raise ExtractionError(f"Failed to write JSON output: {path}") from exc
    except (TypeError, ValueError) as exc:
        raise ExtractionError(f"Failed to serialize report to JSON: {path}") from exc

    logger.info("Saved extraction JSON to %s", path)
    return path.resolve()


def default_json_output_path(docx_path: Path, output_dir: Path) -> Path:
    """Derive ``<stem>.json`` under ``output_dir`` from a DOCX path."""
    return output_dir / f"{Path(docx_path).stem}.json"


def extract_docx_to_json(
    docx_path: Path,
    *,
    output_path: Path | None = None,
    settings: Settings | None = None,
    client: OpenAI | None = None,
) -> Path:
    """
    End-to-end pipeline: DOCX → structured extraction → JSON file.

    Args:
        docx_path: Input ``.docx`` report.
        output_path: Optional explicit JSON path; defaults to ``extracted_json/<stem>.json``.
        settings: Optional runtime settings.
        client: Optional OpenAI client.

    Returns:
        Path to the written JSON file.
    """
    cfg = settings or load_settings()
    source = Path(docx_path).name

    text = load_docx_text(docx_path)
    report = extract_structured_json(
        text,
        source_file=source,
        settings=cfg,
        client=client,
    )

    destination = output_path or default_json_output_path(docx_path, cfg.extracted_json_dir)
    return save_json(report, destination)
